from docx import Document
from nltk.tokenize import word_tokenize
import nltk
import fitz  # PyMuPDF
# nltk.download('punkt')  # 下载必要的数据，只需要运行一次
def count_tokens(paragraph):
    tokens = word_tokenize(paragraph)
    return len(tokens)

def read_paragraphs(file_path):
    doc = Document(file_path)

    paragraphs = []
    for paragraph in doc.paragraphs:
        paragraphs.append(paragraph.text)

    return paragraphs

def group_paragraphs_by_token_limit(file_path, token_limit):
    doc = Document(file_path)

    current_group = []
    total_tokens = 0
    groups = []

    for paragraph in doc.paragraphs:
        tokens_count = count_tokens(paragraph.text)

        if total_tokens + tokens_count > token_limit and current_group:
            # Start a new group when adding the current paragraph exceeds the token limit
            groups.append(current_group)
            current_group = []
            total_tokens = 0

        current_group.append(paragraph.text)
        total_tokens += tokens_count

    # Add the last group if it's not empty
    if current_group:
        groups.append(current_group)

    return groups


def extract_text_from_pages(pdf_path, start_page, end_page):
    # 打开 PDF 文件
    pdf_document = fitz.open(pdf_path)

    # 获取总页数
    total_pages = pdf_document.page_count

    # 确保结束页不超过总页数
    end_page = min(end_page, total_pages)

    # 遍历指定页面范围，提取文本
    extracted_text = ''
    for page_num in range(start_page - 1, end_page):
        page = pdf_document[page_num]
        text = page.get_text()
        extracted_text += f'{text}\n'
        

    # 关闭 PDF 文档
    pdf_document.close()
    return extracted_text

def group_page_by_token_limit(file_path, token_limit,start_page, end_page):
     # 打开 PDF 文件
    pdf_document = fitz.open(file_path)

    # 获取总页数
    total_pages = pdf_document.page_count

    # 确保结束页不超过总页数
    end_page = min(end_page, total_pages)

    current_group = []
    total_tokens = 0
    groups = []

    for page_num in range(start_page - 1, end_page):
        page = pdf_document[page_num]
        tokens_count = count_tokens(page.get_text())
    
        if total_tokens + tokens_count > token_limit and current_group:
            # Start a new group when adding the current paragraph exceeds the token limit
            groups.append(current_group)
            current_group = []
            total_tokens = 0

        current_group.append(page.get_text())
        total_tokens += tokens_count

    # Add the last group if it's not empty
    if current_group:
        groups.append(current_group)

    return groups